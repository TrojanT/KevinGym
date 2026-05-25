from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Open reference doc to inherit styles
doc = Document('/Users/jishananam/Downloads/Assessment 2 - Submission -Group12 - upload.docx')

# Remove all existing content
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)

# ── Colours matching reference doc exactly ─────────────────────────────────────
HEADING_COLOR = RGBColor(0x0F, 0x47, 0x61)   # #0F4761 — doc heading blue
BODY_COLOR    = RGBColor(0x0E, 0x28, 0x41)   # #0E2841 — doc body instruction text
BLACK         = RGBColor(0x00, 0x00, 0x00)
TABLE_HEADER  = '2E75B6'                       # blue from functional testing table
TABLE_ROW_ALT = 'DEEAF1'                       # light blue alt row
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG       = 'F2F2F2'

FONT_BODY    = 'Aptos'
FONT_HEADING = 'Play'

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading1(text):
    """Matches doc Heading 1 — Play 20pt #0F4761"""
    p = doc.add_paragraph(style='Heading 1')
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(20)
    run.font.color.rgb = HEADING_COLOR
    run.bold = False
    return p

def heading2(text):
    """Matches doc Heading 2 — Play 16pt #0F4761"""
    p = doc.add_paragraph(style='Heading 2')
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(16)
    run.font.color.rgb = HEADING_COLOR
    run.bold = False
    return p

def heading3(text):
    """Matches doc Heading 3 — Aptos 14pt #0F4761"""
    p = doc.add_paragraph(style='Heading 3')
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(14)
    run.font.color.rgb = HEADING_COLOR
    run.bold = True
    return p

def body(text):
    """Matches doc body — Aptos 11pt #0E2841"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_COLOR
    return p

def label_line(bold_text, normal_text=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(bold_text)
    r1.font.name = FONT_BODY
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLACK
    r1.bold = True
    if normal_text:
        r2 = p.add_run(normal_text)
        r2.font.name = FONT_BODY
        r2.font.size = Pt(11)
        r2.font.color.rgb = BODY_COLOR
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_COLOR
    return p

def screenshot_note(filename):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'[Insert screenshot: {filename}]')
    run.italic = True
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    return p

def code_block(lines):
    tbl = doc.add_table(rows=1, cols=1)
    _add_tbl_borders(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, CODE_BG)
    # Remove default empty paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)

def _add_tbl_borders(tbl):
    tbl_element = tbl._tbl
    tbl_pr = tbl_element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_element.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'CCCCCC')
        borders.append(b)
    tbl_pr.append(borders)

def info_table(headers, rows, col_widths_cm=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _add_tbl_borders(tbl)
    # Header row — blue bg matching reference (#2E75B6)
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, TABLE_HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        bg = TABLE_ROW_ALT if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = FONT_BODY
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2.1 — DESIGN PATTERNS
# ══════════════════════════════════════════════════════════════════════════════
heading1('2.1  Design Patterns')
body(
    'The KevinGym backend implements five Gang-of-Four design patterns. '
    'Each pattern is described below with the relevant backend source code and justification.'
)

# ── Pattern 1: Singleton ───────────────────────────────────────────────────────
heading2('Pattern 1 — Singleton Pattern')
label_line('Category: ', 'Creational')
label_line('File: ', 'backend/config/db.js')

heading3('Definition')
body(
    'The Singleton Pattern ensures that a class has only one instance throughout the application '
    'lifetime and provides a global point of access to that instance.'
)

heading3('Implementation')
body(
    'The Database class manages the MongoDB connection using the Singleton pattern. It stores a '
    'static instance property initialised to null. The static getInstance() method checks whether '
    'an instance already exists — if not, it creates one; if yes, it returns the existing one. '
    'The connect() method also guards against re-connecting if a live connection already exists.'
)
screenshot_note('backend/config/db.js')
code_block([
    'class Database {',
    '  static instance = null;',
    '',
    '  static getInstance() {',
    '    if (!Database.instance) {',
    '      Database.instance = new Database();',
    '    }',
    '    return Database.instance;',
    '  }',
    '',
    '  async connect() {',
    '    if (this.connection) return;',
    '    this.connection = await mongoose.connect(process.env.MONGO_URI);',
    '    console.log("MongoDB connected successfully");',
    '  }',
    '}',
    '',
    '// server.js — only one connection ever created',
    'Database.getInstance().connect();',
])

heading3('Justification')
body(
    'A database connection is an expensive and limited resource. MongoDB Atlas enforces connection '
    'pool limits, and creating multiple connections wastes memory and can exhaust the pool under load. '
    'The Singleton pattern guarantees that exactly one connection is established at application startup '
    'and reused for all subsequent database operations, ensuring consistency, reliability, and optimal '
    'resource usage across the entire backend.'
)

# ── Pattern 2: Observer ────────────────────────────────────────────────────────
heading2('Pattern 2 — Observer Pattern')
label_line('Category: ', 'Behavioural')
label_line('Files: ', 'backend/observers/LoggerObserver.js  |  backend/observers/NotificationObserver.js  |  backend/events/gymEvents.js')

heading3('Definition')
body(
    'The Observer Pattern defines a one-to-many dependency between objects so that when one object '
    'changes state, all its dependents are notified and updated automatically without tight coupling.'
)

heading3('Implementation')
body(
    'gymEvents (a Node.js EventEmitter) acts as the subject. LoggerObserver and NotificationObserver '
    'subscribe to named events at startup. When key actions occur in a controller — user registered, '
    'course created, membership transitioned, or workout plan assigned — the controller emits a named '
    'event on gymEvents. Both observers react independently: LoggerObserver prints a structured log to '
    'the console, while NotificationObserver persists a notification document to MongoDB.'
)
screenshot_note('backend/observers/LoggerObserver.js')
screenshot_note('backend/observers/NotificationObserver.js')
code_block([
    '// LoggerObserver.js',
    'class LoggerObserver {',
    '  constructor() {',
    '    gymEvents.on("userRegistered", ({ name, email, role }) =>',
    '      console.log(`[LOG] New ${role} registered: ${name} (${email})`)',
    '    );',
    '    gymEvents.on("membershipTransitioned", ({ name, from, to }) =>',
    '      console.log(`[LOG] Membership changed for ${name}: ${from} -> ${to}`)',
    '    );',
    '    gymEvents.on("planCreated", ({ title, memberName }) =>',
    '      console.log(`[LOG] Plan "${title}" assigned to ${memberName}`)',
    '    );',
    '  }',
    '}',
    '',
    '// Emitting from authController.js (the subject)',
    'gymEvents.emit("userRegistered", { name: user.name, email: user.email, role: user.role });',
])

heading3('Justification')
body(
    'The Observer pattern fully decouples business logic (controllers) from side effects (logging, '
    'notifications). Controllers do not import or call LoggerObserver or NotificationObserver directly '
    '— they simply emit an event. Adding a new side effect such as sending an email requires only '
    'writing a new observer class and registering it at startup, with zero changes to any existing '
    'controller code. This follows the Open/Closed Principle and keeps the codebase highly maintainable.'
)

# ── Pattern 3: Strategy ────────────────────────────────────────────────────────
heading2('Pattern 3 — Strategy Pattern')
label_line('Category: ', 'Behavioural')
label_line('Files: ', 'backend/strategies/AuthStrategy.js  |  AuthStrategyContext.js  |  JWTStrategy.js')

heading3('Definition')
body(
    'The Strategy Pattern defines a family of algorithms, encapsulates each one in its own class, '
    'and makes them interchangeable. A context object holds a reference to the current strategy and '
    'delegates work to it without the client needing to know the implementation details.'
)

heading3('Implementation')
body(
    'Authentication is implemented as a strategy. AuthStrategy is the abstract base class defining '
    'a verify(req) interface. JWTStrategy extends it and implements JSON Web Token verification. '
    'AuthStrategyContext stores the active strategy via use() and delegates verify() calls to it. '
    'The protect middleware simply calls AuthStrategyContext.verify(req) without knowing which '
    'authentication mechanism is active. The active strategy is set once at startup in server.js.'
)
screenshot_note('backend/strategies/AuthStrategyContext.js')
screenshot_note('backend/strategies/JWTStrategy.js')
code_block([
    '// AuthStrategyContext.js',
    'class AuthStrategyContext {',
    '  static strategy = null;',
    '  static use(strategy)    { this.strategy = strategy; }',
    '  static async verify(req) {',
    '    if (!this.strategy) throw new Error("No auth strategy set");',
    '    return this.strategy.verify(req);',
    '  }',
    '}',
    '',
    '// JWTStrategy.js',
    'class JWTStrategy extends AuthStrategy {',
    '  async verify(req) {',
    '    const token   = req.headers.authorization.split(" ")[1];',
    '    const decoded = jwt.verify(token, process.env.JWT_SECRET);',
    '    return await User.findById(decoded.id).select("-password");',
    '  }',
    '}',
    '',
    '// server.js — set the active strategy at startup',
    'AuthStrategyContext.use(new JWTStrategy());',
])

heading3('Justification')
body(
    'Authentication methods can evolve over time. A future version of KevinGym might support OAuth, '
    'API keys, or session-based authentication. The Strategy pattern allows any new authentication '
    'method to be introduced by writing one new strategy class and calling AuthStrategyContext.use() '
    'at startup — with zero changes to routes, middleware, or controllers. Each strategy can also '
    'be unit-tested in complete isolation from the rest of the application.'
)

# ── Pattern 4: State ───────────────────────────────────────────────────────────
heading2('Pattern 4 — State Pattern')
label_line('Category: ', 'Behavioural')
label_line('Files: ', 'backend/membership/MembershipState.js  |  MembershipContext.js  |  TrialState.js  |  ActiveState.js  |  SuspendedState.js  |  ExpiredState.js')

heading3('Definition')
body(
    'The State Pattern allows an object to alter its behaviour when its internal state changes, '
    'making it appear to change its class. Each possible state is encapsulated in its own class '
    'rather than handled with large if/else chains in the main object.'
)

heading3('Implementation')
body(
    'Every member has a membership status. MembershipState is the abstract base class defining '
    'the interface: getName(), canBookClass(), canAccessContent(), getDescription(), and '
    'allowedTransitions(). Four concrete states implement this interface. MembershipContext holds '
    'the current state and its transitionTo() method validates that the requested transition is '
    'permitted before switching. Invalid transitions (e.g. Trial directly to Suspended) throw '
    'a descriptive error. Every transition fires a gymEvents event, consumed by both observers.'
)
screenshot_note('backend/membership/MembershipState.js')
screenshot_note('backend/membership/MembershipContext.js')
screenshot_note('backend/membership/ActiveState.js')

body('State transition rules:')
info_table(
    ['State', 'canBookClass', 'canAccessContent', 'Allowed Transitions To'],
    [
        ['Trial',     'Yes', 'Yes (limited)', 'active, expired'],
        ['Active',    'Yes', 'Yes (full)',    'suspended, expired'],
        ['Suspended', 'No',  'No',            'active, expired'],
        ['Expired',   'No',  'No',            'trial, active'],
    ]
)

code_block([
    '// MembershipContext.js',
    'transitionTo(newStatusName) {',
    '  const allowed = this.state.allowedTransitions();',
    '  if (!allowed.includes(newStatusName))',
    '    throw new Error(',
    '      `Cannot transition from "${this.state.getName()}" to "${newStatusName}"`',
    '    );',
    '  this.state = new STATE_MAP[newStatusName]();',
    '}',
    '',
    '// ActiveState.js',
    'class ActiveState extends MembershipState {',
    '  getName()            { return "active"; }',
    '  canBookClass()       { return true; }',
    '  canAccessContent()   { return true; }',
    '  getDescription()     { return "Active membership - full access to all classes."; }',
    '  allowedTransitions() { return ["suspended", "expired"]; }',
    '}',
])

heading3('Justification')
body(
    'Without the State pattern, membership logic would require deeply nested if/else blocks '
    'scattered across controllers (e.g. "if status === suspended, reject booking"). The State '
    'pattern confines each status behaviour to a single self-contained class, making it trivial '
    'to add new states (e.g. FrozenState) without modifying existing code. It also enforces '
    'business rules around state transitions at a single point, preventing invalid jumps '
    'across the entire application.'
)

# ── Pattern 5: Builder ─────────────────────────────────────────────────────────
heading2('Pattern 5 — Builder Pattern')
label_line('Category: ', 'Creational')
label_line('Files: ', 'backend/builder/WorkoutPlanBuilder.js  |  backend/controllers/workoutPlanController.js')

heading3('Definition')
body(
    'The Builder Pattern separates the construction of a complex object from its representation, '
    'allowing the same step-by-step construction process to build different configurations while '
    'validating each property as it is set, before producing the final object via build().'
)

heading3('Implementation')
body(
    'WorkoutPlanBuilder constructs personalised workout plans for gym members. It exposes a fluent '
    'interface where each setter validates its input immediately and returns this to enable method '
    'chaining. The build() method performs a final check before returning the completed plain object '
    'for persistence to MongoDB. The workoutPlanController acts as the director, calling builder '
    'methods in sequence with data from the HTTP request body.'
)
screenshot_note('backend/builder/WorkoutPlanBuilder.js')
screenshot_note('backend/controllers/workoutPlanController.js (the director)')
code_block([
    '// WorkoutPlanBuilder.js',
    'class WorkoutPlanBuilder {',
    '  setTitle(title)        { this._title = title; return this; }',
    '',
    '  setDifficulty(level) {',
    '    const valid = ["beginner", "intermediate", "advanced"];',
    '    if (!valid.includes(level)) throw new Error("Invalid difficulty");',
    '    this._difficulty = level; return this;',
    '  }',
    '',
    '  setDuration(weeks) {',
    '    if (!Number.isInteger(weeks) || weeks < 1)',
    '      throw new Error("Duration must be a positive integer");',
    '    this._durationWeeks = weeks; return this;',
    '  }',
    '',
    '  addExercise({ name, sets, reps, notes = "" }) {',
    '    if (!name) throw new Error("Each exercise must have a name");',
    '    this._exercises.push({ name, sets, reps, notes }); return this;',
    '  }',
    '',
    '  build() {',
    '    if (!this._title) throw new Error("Plan must have a title");',
    '    if (this._exercises.length === 0) throw new Error("Add at least one exercise");',
    '    return { title, difficulty, durationWeeks, exercises, notes };',
    '  }',
    '}',
])
code_block([
    '// workoutPlanController.js — the "director"',
    'const builder = new WorkoutPlanBuilder()',
    '  .setTitle(title)',
    '  .setDifficulty(difficulty)',
    '  .setDuration(parseInt(durationWeeks, 10));',
    '',
    'exercises.forEach(ex => builder.addExercise(ex));',
    'if (notes) builder.setNotes(notes);',
    '',
    'const planData = builder.build();',
    'const plan = await WorkoutPlan.create({ ...planData, memberId, assignedBy });',
])

heading3('Justification')
body(
    'A workout plan has many optional fields: a variable number of exercises, difficulty level, '
    'duration, and notes. Passing all of these as constructor parameters would create an unreadable '
    'signature and make input validation hard to manage. The Builder pattern makes construction '
    'readable through method chaining, enforces validation at each individual step (not just at the '
    'end), and makes it easy to add new plan properties in future without breaking existing '
    'controller code or requiring constructor changes.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2.2 — OOP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('2.2  Implementation of OOP')
body(
    'The KevinGym backend demonstrates all four core Object-Oriented Programming principles across '
    'its class hierarchy. Each principle is identified below with backend code and justification.'
)

# ── OOP 1: Encapsulation ───────────────────────────────────────────────────────
heading2('OOP Principle 1 — Encapsulation')
label_line('Files: ', 'backend/builder/WorkoutPlanBuilder.js  |  backend/config/db.js')

heading3('Definition')
body(
    'Encapsulation is the bundling of data and the methods that operate on that data within a '
    'single class, while restricting direct external access to internal state. It ensures that '
    'an object controls its own integrity.'
)

heading3('Implementation')
body(
    'WorkoutPlanBuilder prefixes all internal properties with an underscore (_title, _exercises, '
    '_difficulty, _durationWeeks, _notes) to signal they are private implementation details. '
    'External code interacts only through the public setter methods, each of which validates input '
    'before modifying internal state. The Database class similarly encapsulates the Mongoose '
    'connection object — callers interact only through getInstance() and connect(), never directly '
    'with this.connection.'
)
screenshot_note('backend/builder/WorkoutPlanBuilder.js — constructor showing private fields')
code_block([
    '// WorkoutPlanBuilder.js — encapsulated private state',
    'constructor() {',
    '  this._title         = "";          // private — only writable via setTitle()',
    '  this._difficulty    = "beginner";  // private — validated via setDifficulty()',
    '  this._durationWeeks = 4;           // private — validated via setDuration()',
    '  this._exercises     = [];          // private — validated via addExercise()',
    '  this._notes         = "";          // private — writable via setNotes()',
    '}',
    '',
    '// Public method controls and validates access',
    'setDifficulty(difficulty) {',
    '  const valid = ["beginner", "intermediate", "advanced"];',
    '  if (!valid.includes(difficulty))',
    '    throw new Error(`Difficulty must be one of: ${valid.join(", ")}`);',
    '  this._difficulty = difficulty;',
    '  return this;',
    '}',
])

heading3('Justification')
body(
    'Encapsulation prevents external code from bypassing validation logic. If _exercises were '
    'publicly accessible, any caller could push a malformed object directly, bypassing the '
    'name-check in addExercise(). By channelling all state changes through validated public '
    'methods, the object guarantees its own internal consistency at all times. This reduces bugs '
    'caused by invalid state and makes each class safe to use across the entire application '
    'without defensive checks in every consumer.'
)

# ── OOP 2: Abstraction ─────────────────────────────────────────────────────────
heading2('OOP Principle 2 — Abstraction')
label_line('Files: ', 'backend/strategies/AuthStrategy.js  |  backend/membership/MembershipState.js')

heading3('Definition')
body(
    'Abstraction hides internal implementation complexity and exposes only what is necessary for '
    'other parts of the system to interact with an object. Abstract base classes define the '
    'interface (the what) without dictating the implementation (the how).'
)

heading3('Implementation')
body(
    'Two abstract base classes are used. AuthStrategy defines the verify(req) interface for all '
    'authentication strategies. MembershipState defines the full interface for all membership '
    'states: getName(), canBookClass(), canAccessContent(), getDescription(), and allowedTransitions(). '
    'Both classes throw descriptive errors if a subclass fails to implement a required method, '
    'enforcing the contract at runtime.'
)
screenshot_note('backend/membership/MembershipState.js')
screenshot_note('backend/strategies/AuthStrategy.js')
code_block([
    '// MembershipState.js — abstract base class',
    'class MembershipState {',
    '  getName()            { throw new Error("getName() not implemented"); }',
    '  canBookClass()       { throw new Error("canBookClass() not implemented"); }',
    '  canAccessContent()   { throw new Error("canAccessContent() not implemented"); }',
    '  getDescription()     { throw new Error("getDescription() not implemented"); }',
    '  allowedTransitions() { throw new Error("allowedTransitions() not implemented"); }',
    '}',
    '',
    '// AuthStrategy.js — abstract base class',
    'class AuthStrategy {',
    '  async verify(req) {',
    '    throw new Error("verify() must be implemented by a concrete strategy");',
    '  }',
    '}',
])

heading3('Justification')
body(
    'Abstraction allows MembershipContext and the protect middleware to operate against a stable '
    'interface without knowing which concrete class is behind it at runtime. The protect middleware '
    'calls AuthStrategyContext.verify(req) without knowing whether it is JWT, OAuth, or any other '
    'mechanism. MembershipContext calls this.state.canBookClass() without knowing which state is '
    'active. This reduces coupling between components and makes the system significantly easier '
    'to extend, test, and maintain independently.'
)

# ── OOP 3: Inheritance ─────────────────────────────────────────────────────────
heading2('OOP Principle 3 — Inheritance')
label_line('Files: ', 'backend/strategies/JWTStrategy.js  |  backend/membership/ActiveState.js  |  backend/membership/SuspendedState.js')

heading3('Definition')
body(
    'Inheritance allows a class to acquire the properties and methods of a parent class, '
    'establishing an "is-a" relationship and enabling interface reuse. Subclasses inherit the '
    'contract defined by the parent and override only the behaviour that differs.'
)

heading3('Implementation')
body(
    'Two inheritance hierarchies exist in the backend. In the Strategy layer, JWTStrategy extends '
    'AuthStrategy and inherits the verify() interface while providing its own concrete JWT-based '
    'implementation. In the State layer, all four concrete state classes (TrialState, ActiveState, '
    'SuspendedState, ExpiredState) extend MembershipState, inheriting the full interface contract '
    'and overriding each method with state-specific behaviour.'
)
screenshot_note('backend/strategies/JWTStrategy.js')
screenshot_note('backend/membership/ActiveState.js')
code_block([
    '// JWTStrategy extends AuthStrategy',
    'class JWTStrategy extends AuthStrategy {',
    '  async verify(req) {',
    '    const token   = req.headers.authorization.split(" ")[1];',
    '    const decoded = jwt.verify(token, process.env.JWT_SECRET);',
    '    return await User.findById(decoded.id).select("-password");',
    '  }',
    '}',
    '',
    '// ActiveState extends MembershipState',
    'class ActiveState extends MembershipState {',
    '  getName()            { return "active"; }',
    '  canBookClass()       { return true; }',
    '  canAccessContent()   { return true; }',
    '  getDescription()     { return "Active membership - full access to all classes."; }',
    '  allowedTransitions() { return ["suspended", "expired"]; }',
    '}',
])

heading3('Justification')
body(
    'Inheritance eliminates the need to redefine the interface contract in every concrete class. '
    'All four state classes are guaranteed to implement the same set of methods because they '
    'inherit from MembershipState. This consistency is what enables MembershipContext to call '
    'this.state.canBookClass() on any state object without knowing its concrete type — which '
    'directly enables the polymorphic behaviour described in Principle 4 below.'
)

# ── OOP 4: Polymorphism ────────────────────────────────────────────────────────
heading2('OOP Principle 4 — Polymorphism')
label_line('Files: ', 'backend/membership/MembershipContext.js  |  backend/observers/LoggerObserver.js  |  backend/observers/NotificationObserver.js')

heading3('Definition')
body(
    'Polymorphism allows objects of different classes to be treated through the same interface, '
    'with each class responding to the same method call in its own way. This eliminates conditional '
    'branching from business logic and makes the system open to extension without modification.'
)

heading3('Implementation')
body(
    'Two forms of polymorphism are demonstrated. First, MembershipContext calls '
    'this.state.canBookClass() on whichever concrete state is active. The same method call returns '
    'true for TrialState and ActiveState, but false for SuspendedState and ExpiredState — with no '
    'if/else logic required in the context. Second, both LoggerObserver and NotificationObserver '
    'respond to the same gymEvents events (e.g. "planCreated") but behave completely differently: '
    'one writes to the console, the other saves to MongoDB.'
)
screenshot_note('backend/membership/MembershipContext.js — canBookClass() delegation')
screenshot_note('backend/observers/LoggerObserver.js and NotificationObserver.js side by side')
code_block([
    '// MembershipContext.js — polymorphic delegation, no if/else needed',
    'canBookClass()     { return this.state.canBookClass(); }',
    'canAccessContent() { return this.state.canAccessContent(); }',
    '',
    '// Runtime behaviour depends on which state is active:',
    '// TrialState.canBookClass()     -> true',
    '// ActiveState.canBookClass()    -> true',
    '// SuspendedState.canBookClass() -> false',
    '// ExpiredState.canBookClass()   -> false',
])
code_block([
    '// Same event — two completely different responses (polymorphism via Observer)',
    '',
    '// LoggerObserver.js',
    'gymEvents.on("planCreated", ({ title, memberName }) =>',
    '  console.log(`[LOG] Plan "${title}" assigned to ${memberName}`)  // logs to console',
    ');',
    '',
    '// NotificationObserver.js',
    'gymEvents.on("planCreated", ({ title, memberName }) =>',
    '  Notification.create({ message: `Plan "${title}" assigned to ${memberName}` })  // saves to DB',
    ');',
])

heading3('Justification')
body(
    'Polymorphism removes if/else branching from core logic. Without it, MembershipContext would '
    'need explicit status checks throughout the codebase: if (status === "suspended") { canBook = false }. '
    'With polymorphism, the correct answer is selected at runtime by the active state object. '
    'This makes controllers shorter and cleaner, and adding a new state (e.g. FrozenState) requires '
    'only a new class — no existing code needs modification. The same principle applies to the '
    'observer system: new reactions to events require only a new observer, not changes to controllers.'
)

# ── Summary tables ─────────────────────────────────────────────────────────────
doc.add_page_break()
heading1('Summary')

heading2('Design Patterns')
info_table(
    ['#', 'Pattern', 'Category', 'Key File', 'Feature Implemented'],
    [
        ['1', 'Singleton', 'Creational',  'backend/config/db.js',                     'Single MongoDB connection instance'],
        ['2', 'Observer',  'Behavioural', 'backend/observers/LoggerObserver.js',       'Event-driven logging and notifications'],
        ['3', 'Strategy',  'Behavioural', 'backend/strategies/AuthStrategyContext.js', 'Swappable authentication mechanism'],
        ['4', 'State',     'Behavioural', 'backend/membership/MembershipContext.js',   'Membership lifecycle management'],
        ['5', 'Builder',   'Creational',  'backend/builder/WorkoutPlanBuilder.js',     'Step-by-step workout plan creation'],
    ]
)

heading2('OOP Principles')
info_table(
    ['Principle',      'Key File(s)',                              'Demonstrated By'],
    [
        ['Encapsulation', 'WorkoutPlanBuilder.js, db.js',             'Private _fields, validated public setters'],
        ['Abstraction',   'MembershipState.js, AuthStrategy.js',      'Abstract base classes with enforced interface contracts'],
        ['Inheritance',   'JWTStrategy.js, ActiveState.js',           'extends across Strategy and State class hierarchies'],
        ['Polymorphism',  'MembershipContext.js, LoggerObserver.js',  'Same method call yields different runtime behaviour per class'],
    ]
)

out = '/Users/jishananam/Desktop/QUT/KevinGym/KevinGym_Section2_DesignPatterns_OOP.docx'
doc.save(out)
print(f'Saved: {out}')
